"""
File Parsers for Different File Formats.

This module provides parsers for various file formats including TXT, PDF, 
Word, Excel, and Markdown files with content extraction capabilities.
"""

import os
import re
import tempfile
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Dict, List, Any, Optional, Union
from datetime import datetime
from io import BytesIO

from src.core.logging import LoggerMixin
from .exceptions import ParsingError, UnsupportedFileTypeError, FileCorruptionError


class BaseParser(ABC, LoggerMixin):
    """
    Abstract base class for file parsers.
    
    Provides common interface for parsing different file formats.
    """
    
    def __init__(self):
        """Initialize the parser."""
        self.supported_extensions: List[str] = []
        self.mime_types: List[str] = []
    
    @abstractmethod
    def parse(self, file_path: Union[str, Path], **kwargs) -> Dict[str, Any]:
        """
        Parse a file and extract its content.
        
        Args:
            file_path: Path to the file to parse
            **kwargs: Additional parsing options
            
        Returns:
            Dictionary containing parsed content and metadata
            
        Raises:
            ParsingError: If parsing fails
        """
        pass
    
    def can_parse(self, file_path: Union[str, Path]) -> bool:
        """
        Check if this parser can handle the given file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            True if parser can handle the file
        """
        file_path = Path(file_path)
        return file_path.suffix.lower() in self.supported_extensions
    
    def _validate_file(self, file_path: Union[str, Path]) -> Path:
        """
        Validate that the file exists and is readable.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Path object
            
        Raises:
            ParsingError: If file is not accessible
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise ParsingError(f"File not found: {file_path}")
        
        if not file_path.is_file():
            raise ParsingError(f"Not a file: {file_path}")
        
        if not os.access(file_path, os.R_OK):
            raise ParsingError(f"File not readable: {file_path}")
        
        return file_path
    
    def _extract_basic_metadata(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract basic file metadata.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dictionary containing basic metadata
        """
        try:
            stat = file_path.stat()
            
            return {
                "file_name": file_path.name,
                "file_path": str(file_path),
                "file_size": stat.st_size,
                "created_at": datetime.fromtimestamp(stat.st_ctime),
                "modified_at": datetime.fromtimestamp(stat.st_mtime),
                "accessed_at": datetime.fromtimestamp(stat.st_atime),
                "file_extension": file_path.suffix.lower(),
                "mime_type": self._get_mime_type(file_path)
            }
        except Exception as e:
            self.logger.error(f"Failed to extract basic metadata: {e}")
            return {}
    
    def _get_mime_type(self, file_path: Path) -> Optional[str]:
        """
        Get MIME type for the file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            MIME type or None
        """
        try:
            import mimetypes
            mime_type, _ = mimetypes.guess_type(str(file_path))
            return mime_type
        except Exception:
            return None


class TextParser(BaseParser):
    """
    Parser for plain text files.
    
    Supports .txt files with various encodings.
    """
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.txt']
        self.mime_types = ['text/plain']
    
    def parse(self, file_path: Union[str, Path], encoding: str = 'utf-8', **kwargs) -> Dict[str, Any]:
        """
        Parse a text file.
        
        Args:
            file_path: Path to the text file
            encoding: Text encoding (default: utf-8)
            **kwargs: Additional options
            
        Returns:
            Dictionary containing parsed content and metadata
        """
        try:
            file_path = self._validate_file(file_path)
            
            # Try to read with specified encoding
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
            except UnicodeDecodeError:
                # Fallback to different encodings
                for fallback_encoding in ['latin-1', 'cp1252', 'utf-16']:
                    try:
                        with open(file_path, 'r', encoding=fallback_encoding) as f:
                            content = f.read()
                        encoding = fallback_encoding
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    raise ParsingError(f"Could not decode file with any encoding: {file_path}")
            
            # Extract basic statistics
            lines = content.splitlines()
            word_count = len(content.split())
            char_count = len(content)
            
            # Extract metadata
            metadata = self._extract_basic_metadata(file_path)
            metadata.update({
                "parser_type": "text",
                "encoding": encoding,
                "line_count": len(lines),
                "word_count": word_count,
                "character_count": char_count,
                "non_empty_lines": len([line for line in lines if line.strip()]),
                "has_content": bool(content.strip())
            })
            
            return {
                "content": content,
                "metadata": metadata,
                "raw_content": content,
                "extracted_at": datetime.now()
            }
            
        except Exception as e:
            self.logger.error(f"Failed to parse text file {file_path}: {e}")
            raise ParsingError(f"Text parsing failed: {e}")


class PDFParser(BaseParser):
    """
    Parser for PDF files.
    
    Uses PyPDF library for content extraction.
    """
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.pdf']
        self.mime_types = ['application/pdf']
    
    def parse(self, file_path: Union[str, Path], **kwargs) -> Dict[str, Any]:
        """
        Parse a PDF file.
        
        Args:
            file_path: Path to the PDF file
            **kwargs: Additional options
            
        Returns:
            Dictionary containing parsed content and metadata
        """
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ParsingError("PyPDF library not installed. Install with: pip install pypdf")
        
        try:
            file_path = self._validate_file(file_path)
            
            # Read PDF
            reader = PdfReader(str(file_path))
            
            # Extract text from all pages
            content_parts = []
            page_contents = []
            
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        content_parts.append(page_text)
                        page_contents.append({
                            "page_number": page_num + 1,
                            "content": page_text,
                            "character_count": len(page_text)
                        })
                except Exception as e:
                    self.logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                    continue
            
            # Combine all content
            full_content = '\n\n'.join(content_parts)
            
            # Extract PDF metadata
            pdf_metadata = {}
            if reader.metadata:
                pdf_metadata = {
                    "title": reader.metadata.get("/Title", ""),
                    "author": reader.metadata.get("/Author", ""),
                    "subject": reader.metadata.get("/Subject", ""),
                    "creator": reader.metadata.get("/Creator", ""),
                    "producer": reader.metadata.get("/Producer", ""),
                    "creation_date": reader.metadata.get("/CreationDate", ""),
                    "modification_date": reader.metadata.get("/ModDate", "")
                }
            
            # Extract basic statistics
            word_count = len(full_content.split())
            char_count = len(full_content)
            
            # Extract metadata
            metadata = self._extract_basic_metadata(file_path)
            metadata.update({
                "parser_type": "pdf",
                "page_count": len(reader.pages),
                "word_count": word_count,
                "character_count": char_count,
                "has_content": bool(full_content.strip()),
                "pdf_metadata": pdf_metadata,
                "pages_with_content": len([p for p in page_contents if p["content"].strip()])
            })
            
            return {
                "content": full_content,
                "metadata": metadata,
                "page_contents": page_contents,
                "raw_content": full_content,
                "extracted_at": datetime.now()
            }
            
        except Exception as e:
            self.logger.error(f"Failed to parse PDF file {file_path}: {e}")
            if "corrupted" in str(e).lower() or "damaged" in str(e).lower():
                raise FileCorruptionError(f"PDF file appears to be corrupted: {e}")
            raise ParsingError(f"PDF parsing failed: {e}")


class WordParser(BaseParser):
    """
    Parser for Microsoft Word documents.
    
    Uses python-docx library for content extraction.
    """
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.docx', '.doc']
        self.mime_types = [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/msword'
        ]
    
    def parse(self, file_path: Union[str, Path], **kwargs) -> Dict[str, Any]:
        """
        Parse a Word document.
        
        Args:
            file_path: Path to the Word document
            **kwargs: Additional options
            
        Returns:
            Dictionary containing parsed content and metadata
        """
        try:
            from docx import Document
        except ImportError:
            raise ParsingError("python-docx library not installed. Install with: pip install python-docx")
        
        try:
            file_path = self._validate_file(file_path)
            
            # Handle .doc files (older format)
            if file_path.suffix.lower() == '.doc':
                # Try to convert .doc to .docx using python-docx2txt as fallback
                try:
                    import docx2txt
                    content = docx2txt.process(str(file_path))
                    
                    # Extract basic statistics
                    word_count = len(content.split())
                    char_count = len(content)
                    
                    # Extract metadata
                    metadata = self._extract_basic_metadata(file_path)
                    metadata.update({
                        "parser_type": "word_legacy",
                        "word_count": word_count,
                        "character_count": char_count,
                        "has_content": bool(content.strip()),
                        "document_format": "doc"
                    })
                    
                    return {
                        "content": content,
                        "metadata": metadata,
                        "raw_content": content,
                        "extracted_at": datetime.now()
                    }
                    
                except ImportError:
                    raise ParsingError("docx2txt library not installed. Install with: pip install docx2txt")
                except Exception as e:
                    raise ParsingError(f"Failed to parse .doc file: {e}")
            
            # Handle .docx files
            document = Document(str(file_path))
            
            # Extract text from paragraphs
            content_parts = []
            paragraph_contents = []
            
            for para_num, paragraph in enumerate(document.paragraphs):
                text = paragraph.text.strip()
                if text:
                    content_parts.append(text)
                    paragraph_contents.append({
                        "paragraph_number": para_num + 1,
                        "content": text,
                        "style": paragraph.style.name if paragraph.style else None
                    })
            
            # Extract text from tables
            table_contents = []
            for table_num, table in enumerate(document.tables):
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_text.append(' | '.join(row_text))
                
                if table_text:
                    table_content = '\n'.join(table_text)
                    content_parts.append(table_content)
                    table_contents.append({
                        "table_number": table_num + 1,
                        "content": table_content,
                        "row_count": len(table_text),
                        "column_count": len(table.columns) if table.rows else 0
                    })
            
            # Combine all content
            full_content = '\n\n'.join(content_parts)
            
            # Extract document properties
            core_properties = document.core_properties
            doc_metadata = {
                "title": core_properties.title or "",
                "author": core_properties.author or "",
                "subject": core_properties.subject or "",
                "keywords": core_properties.keywords or "",
                "comments": core_properties.comments or "",
                "last_modified_by": core_properties.last_modified_by or "",
                "created": core_properties.created.isoformat() if core_properties.created else "",
                "modified": core_properties.modified.isoformat() if core_properties.modified else "",
                "last_printed": core_properties.last_printed.isoformat() if core_properties.last_printed else ""
            }
            
            # Extract basic statistics
            word_count = len(full_content.split())
            char_count = len(full_content)
            
            # Extract metadata
            metadata = self._extract_basic_metadata(file_path)
            metadata.update({
                "parser_type": "word",
                "paragraph_count": len(paragraph_contents),
                "table_count": len(table_contents),
                "word_count": word_count,
                "character_count": char_count,
                "has_content": bool(full_content.strip()),
                "document_metadata": doc_metadata,
                "document_format": "docx"
            })
            
            return {
                "content": full_content,
                "metadata": metadata,
                "paragraph_contents": paragraph_contents,
                "table_contents": table_contents,
                "raw_content": full_content,
                "extracted_at": datetime.now()
            }
            
        except Exception as e:
            self.logger.error(f"Failed to parse Word document {file_path}: {e}")
            if "corrupted" in str(e).lower() or "damaged" in str(e).lower():
                raise FileCorruptionError(f"Word document appears to be corrupted: {e}")
            raise ParsingError(f"Word parsing failed: {e}")


class ExcelParser(BaseParser):
    """
    Parser for Microsoft Excel files.
    
    Uses openpyxl and pandas libraries for content extraction.
    """
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.xlsx', '.xls', '.csv']
        self.mime_types = [
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'application/vnd.ms-excel',
            'text/csv'
        ]
    
    def parse(self, file_path: Union[str, Path], **kwargs) -> Dict[str, Any]:
        """
        Parse an Excel file.
        
        Args:
            file_path: Path to the Excel file
            **kwargs: Additional options
            
        Returns:
            Dictionary containing parsed content and metadata
        """
        try:
            import pandas as pd
        except ImportError:
            raise ParsingError("pandas library not installed. Install with: pip install pandas")
        
        try:
            file_path = self._validate_file(file_path)
            
            # Handle different file formats
            if file_path.suffix.lower() == '.csv':
                return self._parse_csv(file_path, **kwargs)
            elif file_path.suffix.lower() in ['.xlsx', '.xls']:
                return self._parse_excel(file_path, **kwargs)
            else:
                raise UnsupportedFileTypeError(f"Unsupported Excel format: {file_path.suffix}")
                
        except Exception as e:
            self.logger.error(f"Failed to parse Excel file {file_path}: {e}")
            if "corrupted" in str(e).lower() or "damaged" in str(e).lower():
                raise FileCorruptionError(f"Excel file appears to be corrupted: {e}")
            raise ParsingError(f"Excel parsing failed: {e}")
    
    def _parse_csv(self, file_path: Path, **kwargs) -> Dict[str, Any]:
        """Parse CSV file."""
        import pandas as pd
        
        try:
            # Try different encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    df = pd.read_csv(file_path, encoding=encoding)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise ParsingError("Could not decode CSV file with any encoding")
            
            # Convert to text content
            content_parts = []
            
            # Add header
            if not df.empty:
                header = ' | '.join(df.columns.astype(str))
                content_parts.append(header)
                
                # Add data rows
                for _, row in df.iterrows():
                    row_text = ' | '.join(row.astype(str))
                    content_parts.append(row_text)
            
            full_content = '\n'.join(content_parts)
            
            # Extract metadata
            metadata = self._extract_basic_metadata(file_path)
            metadata.update({
                "parser_type": "csv",
                "row_count": len(df),
                "column_count": len(df.columns),
                "columns": df.columns.tolist(),
                "has_content": not df.empty,
                "encoding": encoding
            })
            
            return {
                "content": full_content,
                "metadata": metadata,
                "dataframe": df.to_dict('records'),
                "raw_content": full_content,
                "extracted_at": datetime.now()
            }
            
        except Exception as e:
            raise ParsingError(f"CSV parsing failed: {e}")
    
    def _parse_excel(self, file_path: Path, **kwargs) -> Dict[str, Any]:
        """Parse Excel file."""
        import pandas as pd
        
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            sheet_names = excel_file.sheet_names
            
            content_parts = []
            sheet_contents = []
            
            for sheet_name in sheet_names:
                try:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    
                    if not df.empty:
                        # Convert sheet to text
                        sheet_text_parts = [f"Sheet: {sheet_name}"]
                        
                        # Add header
                        header = ' | '.join(df.columns.astype(str))
                        sheet_text_parts.append(header)
                        
                        # Add data rows
                        for _, row in df.iterrows():
                            row_text = ' | '.join(row.astype(str))
                            sheet_text_parts.append(row_text)
                        
                        sheet_content = '\n'.join(sheet_text_parts)
                        content_parts.append(sheet_content)
                        
                        sheet_contents.append({
                            "sheet_name": sheet_name,
                            "content": sheet_content,
                            "row_count": len(df),
                            "column_count": len(df.columns),
                            "columns": df.columns.tolist(),
                            "has_content": not df.empty
                        })
                
                except Exception as e:
                    self.logger.warning(f"Failed to parse sheet {sheet_name}: {e}")
                    continue
            
            full_content = '\n\n'.join(content_parts)
            
            # Extract metadata
            metadata = self._extract_basic_metadata(file_path)
            metadata.update({
                "parser_type": "excel",
                "sheet_count": len(sheet_names),
                "sheet_names": sheet_names,
                "sheets_with_content": len(sheet_contents),
                "has_content": bool(full_content.strip())
            })
            
            return {
                "content": full_content,
                "metadata": metadata,
                "sheet_contents": sheet_contents,
                "raw_content": full_content,
                "extracted_at": datetime.now()
            }
            
        except Exception as e:
            raise ParsingError(f"Excel parsing failed: {e}")


class MarkdownParser(BaseParser):
    """
    Parser for Markdown files.
    
    Uses markdown library for content extraction and conversion.
    """
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.md', '.markdown']
        self.mime_types = ['text/markdown', 'text/x-markdown']
    
    def parse(self, file_path: Union[str, Path], **kwargs) -> Dict[str, Any]:
        """
        Parse a Markdown file.
        
        Args:
            file_path: Path to the Markdown file
            **kwargs: Additional options
            
        Returns:
            Dictionary containing parsed content and metadata
        """
        try:
            file_path = self._validate_file(file_path)
            
            # Read markdown content
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    markdown_content = f.read()
            except UnicodeDecodeError:
                # Fallback to different encodings
                for encoding in ['latin-1', 'cp1252']:
                    try:
                        with open(file_path, 'r', encoding=encoding) as f:
                            markdown_content = f.read()
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    raise ParsingError("Could not decode Markdown file with any encoding")
            
            # Convert to HTML and extract plain text
            html_content = None
            plain_text = markdown_content
            
            try:
                import markdown
                md = markdown.Markdown(extensions=['extra', 'toc'])
                html_content = md.convert(markdown_content)
                
                # Extract plain text from HTML
                import re
                # Remove HTML tags
                plain_text = re.sub(r'<[^>]+>', '', html_content)
                # Clean up whitespace
                plain_text = re.sub(r'\s+', ' ', plain_text).strip()
                
            except ImportError:
                self.logger.warning("markdown library not installed. Using raw markdown content.")
                plain_text = markdown_content
            
            # Extract markdown structure
            lines = markdown_content.splitlines()
            headers = []
            code_blocks = []
            links = []
            
            in_code_block = False
            for line_num, line in enumerate(lines, 1):
                line = line.strip()
                
                # Track code blocks
                if line.startswith('```'):
                    in_code_block = not in_code_block
                    if in_code_block:
                        code_blocks.append({
                            "line_number": line_num,
                            "language": line[3:].strip() if len(line) > 3 else "",
                            "content": ""
                        })
                elif in_code_block and code_blocks:
                    code_blocks[-1]["content"] += line + '\n'
                
                # Extract headers
                if line.startswith('#') and not in_code_block:
                    level = len(line) - len(line.lstrip('#'))
                    title = line.lstrip('#').strip()
                    headers.append({
                        "level": level,
                        "title": title,
                        "line_number": line_num
                    })
                
                # Extract links
                if not in_code_block:
                    link_pattern = r'\[([^\]]+)\]\(([^)]+)\)'
                    matches = re.findall(link_pattern, line)
                    for text, url in matches:
                        links.append({
                            "text": text,
                            "url": url,
                            "line_number": line_num
                        })
            
            # Extract basic statistics
            word_count = len(plain_text.split())
            char_count = len(plain_text)
            
            # Extract metadata
            metadata = self._extract_basic_metadata(file_path)
            metadata.update({
                "parser_type": "markdown",
                "line_count": len(lines),
                "word_count": word_count,
                "character_count": char_count,
                "has_content": bool(plain_text.strip()),
                "header_count": len(headers),
                "code_block_count": len(code_blocks),
                "link_count": len(links),
                "has_html_conversion": html_content is not None
            })
            
            return {
                "content": plain_text,
                "metadata": metadata,
                "markdown_content": markdown_content,
                "html_content": html_content,
                "structure": {
                    "headers": headers,
                    "code_blocks": code_blocks,
                    "links": links
                },
                "raw_content": markdown_content,
                "extracted_at": datetime.now()
            }
            
        except Exception as e:
            self.logger.error(f"Failed to parse Markdown file {file_path}: {e}")
            raise ParsingError(f"Markdown parsing failed: {e}")


class ParserFactory:
    """
    Factory class for creating appropriate parsers based on file types.
    """
    
    def __init__(self):
        """Initialize the parser factory."""
        self.parsers = {
            'text': TextParser(),
            'pdf': PDFParser(),
            'word': WordParser(),
            'excel': ExcelParser(),
            'markdown': MarkdownParser()
        }
    
    def get_parser(self, file_path: Union[str, Path]) -> Optional[BaseParser]:
        """
        Get appropriate parser for the given file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Parser instance or None if no suitable parser found
        """
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        # Check each parser
        for parser in self.parsers.values():
            if parser.can_parse(file_path):
                return parser
        
        return None
    
    def get_supported_extensions(self) -> List[str]:
        """
        Get list of all supported file extensions.
        
        Returns:
            List of supported extensions
        """
        extensions = []
        for parser in self.parsers.values():
            extensions.extend(parser.supported_extensions)
        return sorted(list(set(extensions)))
    
    def is_supported(self, file_path: Union[str, Path]) -> bool:
        """
        Check if the file type is supported.
        
        Args:
            file_path: Path to the file
            
        Returns:
            True if file type is supported
        """
        return self.get_parser(file_path) is not None